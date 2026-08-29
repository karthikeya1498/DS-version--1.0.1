"""Validate the generated OPTIMA-X project report DOCX.

Author: Karthikeya
"""
from pathlib import Path
from docx import Document

report = Path(__file__).resolve().parents[1] / "docs/reports/OPTIMA_X_Complete_Project_Working_Report.docx"
document = Document(report)
words = sum(len(paragraph.text.split()) for paragraph in document.paragraphs)
words += sum(len(cell.text.split()) for table in document.tables for row in table.rows for cell in row.cells)
headings = sum(1 for paragraph in document.paragraphs if paragraph.style.name.startswith("Heading"))
print({"file": str(report), "paragraphs": len(document.paragraphs), "headings": headings, "tables": len(document.tables), "words": words, "bytes": report.stat().st_size})
if words < 7000:
    raise SystemExit("report is shorter than 7,000 words")
if headings < 20:
    raise SystemExit("report is missing expected section headings")
