"""Build the complete OPTIMA-X project report DOCX.

Author: Karthikeya
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/reports/optimax_complete_project_report.md"
OUTPUT = ROOT / "docs/reports/OPTIMA_X_Complete_Project_Working_Report.docx"


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[index], value, bold=True)
        shade_cell(table.rows[0].cells[index], "17324D")
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    document.add_paragraph()


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = "No Spacing"
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.right_indent = Inches(0.35)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        if index < len(lines) - 1:
            run.add_break()
    paragraph.paragraph_format.space_after = Pt(8)


def add_inline_markdown(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|> .*)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            run = paragraph.add_run(token)
            run.italic = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def configure_document(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in (("Title", 24, "17324D"), ("Heading 1", 16, "17324D"), ("Heading 2", 13, "267A78"), ("Heading 3", 11, "17324D")):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    header = section.header.paragraphs[0]
    header.text = "OPTIMA-X | Complete Project Working Report"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(90, 105, 120)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Karthikeya  •  OPTIMA-X  •  Phase 1–7 Review")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(90, 105, 120)


def build() -> None:
    document = Document()
    configure_document(document)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        if line.startswith("```"):
            code = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                code.append(lines[index])
                index += 1
            add_code_block(document, code)
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                values = [value.strip() for value in lines[index].strip().strip("|").split("|")]
                if not all(set(value) <= {"-", ":", " "} for value in values):
                    table_lines.append(values)
                index += 1
            add_table(document, table_lines)
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(line[2:].strip())
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
        elif line.startswith("> "):
            paragraph = document.add_paragraph(style="Intense Quote")
            add_inline_markdown(paragraph, line[2:].strip())
        elif line.startswith("---"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run("─" * 92)
            run.font.color.rgb = RGBColor(180, 190, 200)
        elif re.match(r"^\d+\. ", line):
            paragraph = document.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, line[2:].strip())
        else:
            paragraph = document.add_paragraph()
            add_inline_markdown(paragraph, line.strip())
        index += 1
    document.core_properties.title = "OPTIMA-X Complete Project Working Report"
    document.core_properties.author = "Karthikeya"
    document.core_properties.subject = "Phase 1–7 architecture, implementation, validation, and roadmap"
    document.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
